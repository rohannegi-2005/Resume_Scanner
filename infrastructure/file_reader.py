"""
FileReader v2 — Handles both normal and space-stripped PDFs

The core problem with many LaTeX/Word-exported PDFs:
  pdfplumber's default extract_text() concatenates characters without
  spaces because the font's kerning data has zero-width space glyphs.

Solution: char-level extraction — measure the horizontal gap between
consecutive characters on each line. If gap > 15% of font size → insert
a space. This correctly reconstructs words in ALL tested resume templates
without breaking normally-spaced PDFs.

Fallback: if no chars metadata (rare, image-based PDFs), falls back to
raw extract_text() result.
"""

import os
from collections import defaultdict

import pdfplumber
from docx import Document


class FileReader:
    SUPPORTED = {".pdf", ".docx"}

    def read(self, file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return self._read_pdf(file_path)
        elif ext == ".docx":
            return self._read_docx(file_path)
        raise ValueError(f"Unsupported file type: {ext}")

    # ─────────────────────────────────────────────────────────────────
    # PDF extraction — char-gap approach
    # ─────────────────────────────────────────────────────────────────
    def _read_pdf(self, path: str) -> str:
        pages_text = []
        try:
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    text = self._extract_page(page)
                    if text:
                        pages_text.append(text)
        except Exception as e:
            raise RuntimeError(f"PDF read error: {e}")

        if not pages_text:
            raise RuntimeError("PDF appears empty or image-only (no extractable text).")
        return "\n".join(pages_text)

    def _extract_page(self, page) -> str:
        """
        Char-gap reconstruction:
          1. Group characters by their vertical position (2pt buckets).
          2. Sort each row left-to-right by x0.
          3. Insert a space between chars when gap > 15% of avg font size.
          4. Skip control chars and empty chars.
        """
        chars = page.chars
        if not chars:
            # Fallback for image-based pages
            return page.extract_text() or ""

        lines: dict = defaultdict(list)
        for ch in chars:
            if not ch.get("text", "").strip():
                continue                         # skip whitespace/control chars
            key = round(float(ch["top"]) / 2) * 2
            lines[key].append(ch)

        text_lines = []
        for top_key in sorted(lines.keys()):
            row = sorted(lines[top_key], key=lambda c: float(c["x0"]))
            reconstructed = ""
            prev = None
            for ch in row:
                if prev is not None:
                    gap = float(ch["x0"]) - float(prev["x1"])
                    avg_size = (float(ch.get("size", 10)) + float(prev.get("size", 10))) / 2
                    # 15% of font size is a robust threshold:
                    #   - normal letter-spacing gaps ≈ 0-5%  → no space
                    #   - inter-word gaps ≈ 20-40%          → space
                    if gap > avg_size * 0.15:
                        reconstructed += " "
                reconstructed += ch["text"]
                prev = ch

            stripped = reconstructed.strip()
            if stripped:
                text_lines.append(stripped)

        return "\n".join(text_lines)

    # ─────────────────────────────────────────────────────────────────
    # DOCX extraction
    # ─────────────────────────────────────────────────────────────────
    def _read_docx(self, path: str) -> str:
        try:
            doc = Document(path)
            parts = []
            # Paragraphs (main body)
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            # Tables (skills grids etc.)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text.strip())
            return "\n".join(parts).strip()
        except Exception as e:
            raise RuntimeError(f"DOCX read error: {e}")
